# -*- coding: utf-8 -*-
"""Genera el documento del gatillo de pullback direccional."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DEST = (r"C:\Users\gcarv\Documents\CARPETA PERSONAL\01. guillermo carvajal"
        r"\01_Sigma\mentoria alejandro\estrategias automatizadas"
        r"\estrategia direccional\mejoras al algoritmo direccional")
ARCHIVO = "Gatillo de Pullback - Estrategia Direccional SPX.docx"

AZUL = RGBColor(0x1F, 0x38, 0x64)
GRIS = RGBColor(0x59, 0x59, 0x59)
ROJO = RGBColor(0xC0, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.left_margin = s.right_margin = Inches(0.9)

est = d.styles["Normal"]
est.font.name = "Calibri"
est.font.size = Pt(10.5)


def h(txt, nivel=1):
    p = d.add_heading(txt, level=nivel)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p


def par(txt, negrita=False, color=None, size=None, espacio=6):
    p = d.add_paragraph()
    r = p.add_run(txt)
    r.bold = negrita
    if color: r.font.color.rgb = color
    if size:  r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(espacio)
    return p


def code(lineas):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("\n".join(lineas))
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def tabla(cabeceras, filas, anchos=None):
    t = d.add_table(rows=1, cols=len(cabeceras))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cabeceras):
        cel = t.rows[0].cells[i]
        cel.text = ""
        r = cel.paragraphs[0].add_run(c)
        r.bold = True
        r.font.size = Pt(9.5)
    for fila in filas:
        cells = t.add_row().cells
        for i, v in enumerate(fila):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
    if anchos:
        for fila in t.rows:
            for i, w in enumerate(anchos):
                fila.cells[i].width = Inches(w)
    d.add_paragraph()
    return t


# ── Portada ────────────────────────────────────────────────────────────────
tit = d.add_heading("Gatillo de Pullback", level=0)
for r in tit.runs:
    r.font.color.rgb = AZUL
par("Estrategia Direccional SPX 0DTE — documento de diseño y validación", size=12, color=GRIS)
par("Versión implementada el 3 de agosto de 2026 · activa en producción desde el 4 de agosto",
    size=9.5, color=GRIS, espacio=2)
par("Configuración: spxConfig.entryMode = 'pullback' · reversible a 'camino_b' sin desplegar",
    size=9.5, color=GRIS, espacio=18)

# ── Resumen ────────────────────────────────────────────────────────────────
h("En una frase", 1)
par("El sistema ya no entra cuando la tendencia de 2 minutos está activa —con el precio ya "
    "extendido— sino cuando el precio VUELVE a la EMA10 y reanuda, que es donde el recorrido "
    "todavía está por delante.")

h("La estructura: dos capas", 1)
par("Capa 1 — El marco maestro (15 minutos).", negrita=True, espacio=2)
par("Decide si se puede operar y en qué dirección. Nada entra si esto no se cumple:")
code(["Fase 2 (alcista):   precio > EMA20  ·  EMA10 > EMA20  ·  EMA20 subiendo",
      "Fase 4 (bajista):   precio < EMA20  ·  EMA10 < EMA20  ·  EMA20 bajando"])
par("Si no hay ni Fase 2 ni Fase 4, ni se evalúa el pullback. Este marco también fija la "
    "dirección: no hay decisión de comprar o vender en 2 minutos.")

par("Capa 2 — El gatillo (2 minutos).", negrita=True, espacio=2)
par("Ya sabiendo la dirección, busca el momento de entrar.")

h("La medida clave: distancia a la EMA10 con signo a favor", 1)
code(["d = dir > 0 ? (close − EMA10)    // alcista: positivo = precio ARRIBA de la EMA10",
      "            : (EMA10 − close)    // bajista: positivo = precio ABAJO de la EMA10"])
par("Esto hace que la lógica sea una sola para alcista y bajista. En ambos casos:")
tabla(["Valor de d", "Significado"],
      [["Grande y positivo", "Precio extendido, el movimiento ya corrió"],
       ["Chico", "Precio pegado a la EMA10"],
       ["Negativo", "Precio cruzó al otro lado (el retroceso pasó de largo)"]],
      [1.6, 4.6])

h("Los dos disparadores", 1)
par("(a) CRUCE — el precio estaba del otro lado y vuelve a cruzar la EMA10", negrita=True, espacio=2)
code(["alcista:  close[anterior] < EMA10   Y   close[actual] > EMA10",
      "bajista:  close[anterior] > EMA10   Y   close[actual] < EMA10"])
par("Es la «Conservative Entry» del script original de Alejandro (entryUpT_SS). Retroceso "
    "profundo que se agota y reanuda.")

par("(b) ROCE — el precio se acercó, no llegó a cruzar, y giró", negrita=True, espacio=2)
code(["giro  = d[actual] > d[anterior]   Y   d[anterior] ≤ d[hace 2]",
      "cerca = |d[anterior]| ≤ 0.8 × ATR(14)",
      "roce  = giro Y cerca"])
par("Hizo falta porque dos de los cinco casos identificados nunca cruzaron la EMA10 "
    "(mínimos de +0.31 y +0.66 ATR). Con solo el cruce se habrían perdido.")
par("Dispara si cualquiera de los dos se cumple.", negrita=True)

h("Por qué el umbral va en ATR y no en puntos", 1)
par("Los cinco casos identificados tuvieron profundidades muy distintas:")
tabla(["Caso", "Profundidad", "Observación"],
      [["3-ago 10:06", "+0.66 ATR", "ni tocó la EMA10"],
       ["3-ago 10:32", "+0.31 ATR", "ni tocó"],
       ["3-ago 13:06", "−0.66 ATR", "cruzó apenas"],
       ["9-jul 10:20", "−2.00 ATR", "cruzó profundo"],
       ["13-jul 11:36", "−0.30", "tocó justo (bajista)"]],
      [1.5, 1.5, 3.2])
par("En puntos eso va de 1 a 12. Un umbral fijo en puntos serviría para una volatilidad y no "
    "para otra. En ATR, «cerca» significa lo mismo a las 9:45 con el mercado nervioso que a "
    "la 1 de la tarde con todo quieto.")

d.add_page_break()

h("Ejemplo real paso a paso — el caso bajista del 13 de julio", 1)
code(["hora    close     EMA10     d      qué pasa",
      "11:30   7549.82   7548.89  -0.93   precio ARRIBA de la EMA10 (retroceso en curso)",
      "11:32   7549.31   7548.96  -0.35   sigue arriba",
      "11:36   7548.72   7548.42  -0.30   todavía arriba   ->  NO dispara",
      "11:38   7547.30   7548.21  +0.91   cruza de vuelta  ->  DISPARA"])
par("A las 11:36 el precio seguía dentro del retroceso. A las 11:38 reanuda, y ahí entra. "
    "Ese punto dio +29.8 puntos a favor sin una sola vela en contra.")
par("Detalle relevante: la entrada identificada visualmente fue 11:50. El gatillo entra 12 "
    "minutos antes y 10 puntos mejor — no porque sea más listo, sino porque no espera a que "
    "el movimiento sea obvio.", negrita=True)

h("Qué se quitó, y por qué", 1)
tabla(["Filtro", "Estado", "Motivo"],
      [["MACD de 2m", "FUERA",
        "Un retroceso hunde la línea del MACD contra su señal por construcción, así que vetaba "
        "justo el momento buscado. Y el 9-jul estuvo A FAVOR en las tres peores señales del día "
        "(−30.8, −27.4, −8.3 pts de excursión adversa), mientras las señales con el MACD en "
        "contra tuvieron mediana de −0.5. No filtra riesgo: hace entrar tarde."],
       ["EMAs alineadas en 2m", "FUERA",
        "El pullback ES una fase contraria breve dentro de la fase mayor. El 9-jul la nube de 2m "
        "se fue a −3.58 (Fase 3 pura) y ahí estaba la mejor entrada del día: +25 pts en 18 min."],
       ["Trend Magic", "FUERA", "No formaba parte de la variante medida."],
       ["Marco 15m", "SE QUEDA",
        "Se cumplió en los cinco casos identificados. Es el filtro que sí hace su trabajo."]],
      [1.5, 0.9, 3.8])

h("Qué NO cambió", 1)
par("Después del gatillo, todo el resto del pipeline sigue igual: score del playbook ≥80 "
    "(o ≥90 si el trade anterior del día cerró en pérdida), selección de strikes por delta "
    "0.30, ancho de 10 puntos, TP 30% / SL 50%, stop técnico, exclusividad de posición y "
    "ventana de 9:45am a 2:00pm ET.")
par("El pullback cambia CUÁNDO se dispara la evaluación, no qué se opera ni cómo se gestiona.",
    negrita=True)

h("Evidencia que respalda el cambio", 1)
par("Estudio sobre 18 días de velas reales (9-jul a 3-ago), mismo simulador para las tres "
    "variantes, con exclusividad de posición simulada:")
tabla(["Prueba", "Resultado"],
      [["7 combinaciones de TP/SL", "El pullback gana en 6 de 7"],
       ["Muestra partida en dos mitades", "Gana en AMBAS por separado (9-21 jul y 22 jul-3 ago)"],
       ["Pullback + MACD", "Peor que pullback solo en 6 de 7 combinaciones"],
       ["Los 5 casos identificados", "Los 5 disparan, incluido el bajista"]],
      [2.4, 3.8])
par("Que gane en las dos mitades por separado es lo que descarta el sobreajuste: si la ventaja "
    "solo apareciera en el agregado, sería sospechosa.")

h("Limitaciones — leer antes de sacar conclusiones", 1)
par("1. La ventaja está medida ENTRE VARIANTES DE SIMULACIÓN, no contra P&L real.", negrita=True, espacio=2)
par("El histórico real está contaminado: 39 de los 62 trades direccionales tienen el P&L mal "
    "calculado por el método /gainloss anterior, que asignaba mal las patas cuando varios "
    "trades compartían strikes el mismo día. Al menos 5 registran una pérdida MAYOR al débito "
    "pagado, algo físicamente imposible en un spread de débito. Esos registros son "
    "irrecuperables.", color=ROJO)
par("2. El comportamiento bajista nunca se vio en vivo.", negrita=True, espacio=2)
par("Los cinco casos validados son reconstrucción histórica.")
par("3. El score puede estrangular la mejora.", negrita=True, espacio=2)
par("El gatillo dispara más seguido, pero después hay que pasar el 80%, donde "
    "patrones_estructurales y macd_cruce_pendiente son casi obligatorios en la práctica.")
par("Conclusión honesta: el pullback debería ser mejor que la lógica anterior. No hay evidencia "
    "suficiente para afirmar que sea rentable.", negrita=True, color=ROJO)

h("Qué vigilar en los primeros días", 1)
tabla(["Qué", "Referencia", "Qué haría si se desvía"],
      [["Señales por día", "5-6 esperadas (antes 3.9)",
        "Muy por encima: apretar el 0.8 × ATR"],
       ["Bajista en vivo", "Nunca observado", "Revisar la primera que dispare, vela por vela"],
       ["Rechazos por score", "Desconocido",
        "Si el gatillo mejora y el score se come la mejora, revisar los pesos"]],
      [1.6, 2.0, 2.6])

h("Cómo revertir", 1)
code(['POST /api/spx/config   {"entryMode": "camino_b"}'])
par("Vuelve todo al comportamiento anterior sin tocar código ni desplegar.")

h("Referencias técnicas", 1)
tabla(["Elemento", "Ubicación"],
      [["Función del gatillo", "src/camino_b.js → calcPullbackEntry()"],
       ["Marco maestro 15m", "src/camino_b.js → calcFase15mSimple()"],
       ["Conexión", "server.js → checkDirectionalAutonomous() (cada 30s)"],
       ["Interruptor", "spxConfig.entryMode"],
       ["Umbral tras pérdida", "spxConfig.minScoreTrasPerdida = 90"],
       ["Control de cambios", "control_cambios_direccional.xlsx"],
       ["Resultados por versión", "GET /api/spx/version-stats"]],
      [2.2, 4.0])

os.makedirs(DEST, exist_ok=True)
salida = os.path.join(DEST, ARCHIVO)
try:
    d.save(salida)
    print("guardado:", salida)
except PermissionError:
    alt = salida.replace(".docx", "_nuevo.docx")
    d.save(alt)
    print("ABIERTO EN WORD, guardado como:", alt)
