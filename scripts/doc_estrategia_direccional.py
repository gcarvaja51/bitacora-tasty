# -*- coding: utf-8 -*-
"""Resumen de la estrategia direccional: antes vs ahora, produccion vs monitoreo."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

DEST = (r"C:\Users\gcarv\Documents\CARPETA PERSONAL\01. guillermo carvajal"
        r"\01_Sigma\mentoria alejandro\estrategias automatizadas"
        r"\estrategia direccional\mejoras al algoritmo direccional")
ARCHIVO = "Estrategia Direccional SPX - Antes y Ahora.docx"

AZUL  = RGBColor(0x1F, 0x38, 0x64)
GRIS  = RGBColor(0x59, 0x59, 0x59)
ROJO  = RGBColor(0xC0, 0x00, 0x00)
VERDE = RGBColor(0x1E, 0x6B, 0x3A)

d = Document()
for s in d.sections:
    s.left_margin = s.right_margin = Inches(0.85)
d.styles["Normal"].font.name = "Calibri"
d.styles["Normal"].font.size = Pt(10.5)


def h(t, n=1):
    p = d.add_heading(t, level=n)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p


def par(t, b=False, c=None, sz=None, sp=6):
    p = d.add_paragraph()
    r = p.add_run(t)
    r.bold = b
    if c: r.font.color.rgb = c
    if sz: r.font.size = Pt(sz)
    p.paragraph_format.space_after = Pt(sp)
    return p


def bullet(t, c=None):
    p = d.add_paragraph(style="List Bullet")
    r = p.add_run(t)
    if c: r.font.color.rgb = c
    p.paragraph_format.space_after = Pt(3)
    return p


def code(l):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("\n".join(l))
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def tabla(cab, filas, anchos=None, estilo="Light Grid Accent 1"):
    t = d.add_table(rows=1, cols=len(cab))
    t.style = estilo
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cab):
        cel = t.rows[0].cells[i]
        cel.text = ""
        r = cel.paragraphs[0].add_run(c)
        r.bold = True
        r.font.size = Pt(9.5)
    for f in filas:
        cs = t.add_row().cells
        for i, v in enumerate(f):
            cs[i].text = ""
            r = cs[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
    if anchos:
        for fila in t.rows:
            for i, w in enumerate(anchos):
                fila.cells[i].width = Inches(w)
    d.add_paragraph()
    return t


# ── Portada ────────────────────────────────────────────────────────────────
tit = d.add_heading("Estrategia Direccional SPX 0DTE", level=0)
for r in tit.runs:
    r.font.color.rgb = AZUL
par("Cómo se gestionaba antes y cómo se gestiona desde hoy", sz=12, c=GRIS)
par("Actualizado el 4 de agosto de 2026", sz=9.5, c=GRIS, sp=16)

h("Resumen en tres frases", 1)
par("1. La ENTRADA cambió: antes se entraba cuando la tendencia de 2 minutos ya estaba "
    "activa —con el precio extendido—, ahora se entra cuando el precio vuelve a la EMA10 "
    "y reanuda.")
par("2. La SALIDA no cambió todavía, pero se descubrió que es el problema más grande: "
    "el TP del 30% cierra en ~9 puntos de SPX un movimiento que mediana 44.")
par("3. Se agregó un freno: tras una pérdida del día, el siguiente trade exige 90% de "
    "score en vez de 80%.")

d.add_page_break()

# ── ANTES / AHORA ──────────────────────────────────────────────────────────
h("Antes y ahora, lado a lado", 1)
tabla(["", "ANTES", "DESDE HOY"],
      [["Marco maestro 15m", "Fase 2 o Fase 4 (obligatorio)", "Igual — sin cambios"],
       ["Gatillo en 2m",
        "Alineación activa: precio sobre Trend Magic + EMA10>EMA20 + MACD línea>señal",
        "PULLBACK: el precio vuelve a la EMA10 y reanuda"],
       ["MACD de 2m", "Obligatorio", "ELIMINADO del gatillo"],
       ["EMAs alineadas en 2m", "Obligatorio", "ELIMINADO del gatillo"],
       ["Trend Magic", "Obligatorio", "ELIMINADO del gatillo"],
       ["Score mínimo", "80%", "80%, y 90% si el trade anterior del día perdió"],
       ["Ventana horaria", "9:45am – 2:30pm ET", "9:45am – 2:00pm ET"],
       ["Take Profit", "30% del débito pagado", "Igual — en revisión"],
       ["Stop Loss", "50% del débito", "Igual"],
       ["Stop técnico", "Fractal 15m / POC", "Igual — bajo sospecha"],
       ["Temporalidad", "2 minutos", "2 minutos (el 1m quedó descartado)"]],
      [1.5, 2.6, 2.9])

h("El gatillo nuevo, en detalle", 1)
par("Capa 1 — Marco maestro (15 minutos). Define si se opera y en qué dirección:", b=True, sp=2)
code(["Fase 2 (alcista):  precio > EMA20  ·  EMA10 > EMA20  ·  EMA20 subiendo",
      "Fase 4 (bajista):  precio < EMA20  ·  EMA10 < EMA20  ·  EMA20 bajando"])
par("Si no hay ni Fase 2 ni Fase 4, no se evalúa nada.")

par("Capa 2 — Gatillo (2 minutos). Dispara de dos formas:", b=True, sp=2)
code(["(a) CRUCE   el precio cerró del otro lado de la EMA10 y vuelve a cruzarla",
      "(b) ROCE    la distancia a la EMA10 tocó un mínimo local y giró,",
      "            con ese mínimo dentro de 0.8 × ATR(14) de 2m"])
par("El umbral va en ATR y no en puntos porque la profundidad útil de un retroceso cambia "
    "con la volatilidad: en los casos reales fue de +0.66 a −2.0 ATR, o sea de 1 a 12 puntos.")

h("Por qué se sacó el MACD", 1)
par("Es el cambio que más puede sorprender, así que conviene tener la razón a mano:")
bullet("Un retroceso hunde la línea del MACD contra su señal POR CONSTRUCCIÓN. El filtro "
       "vetaba exactamente el momento que la estrategia busca.")
bullet("Caso real del 3-ago 10:34: Trend Magic, EMAs y marco de 15m en orden; el MACD en "
       "contra las 8 velas del pullback. El sistema no entró; el movimiento salió.")
bullet("Caso real del 9-jul: el MACD dio luz verde 14 minutos y 21 puntos DESPUÉS del piso.")
bullet("Y no protege: ese mismo 9-jul estuvo A FAVOR en las TRES peores señales del día "
       "(−30.8, −27.4 y −8.3 puntos de excursión adversa), mientras las señales con el MACD "
       "en contra tuvieron una mediana de −0.5.")

d.add_page_break()

h("Validación contra las operaciones del mentor", 1)
par("Se reconstruyeron 14 fotos de operaciones exitosas de Alejandro (15-jun a 13-jul) con "
    "velas reales de Tradier. Tres resultaron ser de Reversión a la Media y se descartaron; "
    "quedaron 11 direccionales, de las cuales 8 se pudieron validar con datos.")
tabla(["Fecha", "Entrada del mentor", "Nuestro gatillo", "Diferencia"],
      [["23-jun", "Corto @ 7395", "10:52  S@7398", "+3 pts"],
       ["24-jun", "Corto @ 7390", "13:04  S@7389", "−1 pt"],
       ["30-jun", "Largo @ 7462", "10:42  L@7472", "+10 pts"],
       ["30-jun", "Largo @ 7475", "11:40  L@7480", "+5 pts"],
       ["02-jul", "Corto @ 7490", "11:42  S@7483", "−7 pts"],
       ["09-jul", "Largo @ 7503", "10:24  L@7497", "−6 pts"],
       ["09-jul", "Largo @ 7532", "12:48  L@7532", "0 pts"],
       ["13-jul", "Corto @ 7548", "11:50  S@7539", "−9 pts"]],
      [1.0, 1.9, 2.0, 1.2])
par("8 de 8 en dirección y momento, con mediana de 4 puntos de diferencia. En seis de ocho "
    "la entrada nuestra fue igual o mejor que la suya.", b=True)
par("En las 11 direccionales el marco de 15m estaba alineado con la dirección del trade, "
    "sin una sola excepción. Eso valida el gate maestro.")

h("El hallazgo más importante, y todavía sin resolver", 1)
par("La salida está mal dimensionada frente al movimiento real:", b=True)
code(["TP actual (+30% del débito)  se alcanza con   ~ +9 puntos de SPX",
      "SL actual (−50% del débito)  se alcanza con   ~ −16 puntos de SPX",
      "Recorrido mediano tras las entradas del mentor:  +43.9 puntos"])
par("Se cierra en una quinta parte del movimiento, arriesgando 16 puntos para ganar 9. Con "
    "esa geometría hay que acertar 64% solo para empatar, y 9 puntos de SPX en 0DTE se "
    "cruzan por ruido.", c=ROJO)
par("Comparación sobre las MISMAS 9 entradas del mentor, cambiando solo la salida:")
tabla(["Regla de salida", "Puntos por trade"],
      [["TP actual (+30%)", "9.0"],
       ["Cierre cruza la EMA10 en contra", "17.1"],
       ["Cede 30% del pico", "18.8"],
       ["Cede 1.5 × ATR desde el pico", "19.2"],
       ["Cierre cruza la EMA20 en contra", "21.6"],
       ["Cede 50% del pico", "22.7"],
       ["Cede 2.5 × ATR desde el pico  ← la elegida", "23.7"],
       ["(sin trailing, aguantar al cierre)", "31.3"]],
      [4.2, 1.9])
par("Dato contraintuitivo: usar la EMA10 como piso resultó la regla MÁS floja. Tiene "
    "sentido — es justo la línea que el gatillo de entrada considera una visita sana del "
    "precio, así que usarla de salida sería cerrar en cada retroceso normal.")

d.add_page_break()

# ── PRODUCCION ─────────────────────────────────────────────────────────────
h("Lo que YA está en producción", 1)
par("Activo y operando con dinero (cuenta demo de Tradier) desde el 4 de agosto.", c=VERDE, b=True)
tabla(["Qué", "Detalle", "Cómo revertir"],
      [["Gatillo de pullback", "entryMode = 'pullback'",
        "POST /api/spx/config  {\"entryMode\": \"camino_b\"}"],
       ["Umbral 90% tras pérdida", "minScoreTrasPerdida = 90",
        "Poner el mismo valor que minScore (80)"],
       ["Ventana hasta las 2:00pm ET", "Antes 2:30pm", "Requiere cambio de código"],
       ["Sello de versión por trade", "algoVersion en cada ejecución", "—"],
       ["Redondeo de precios a 2 decimales", "Y en la dirección correcta según abrir/cerrar", "—"],
       ["Reconciliación de P&L corregida", "Usa la orden que realmente cerró", "—"]],
      [1.9, 2.4, 2.7])

h("Lo que está en MONITOREO (registra, no ejecuta)", 1)
par("Corre en paralelo y anota qué habría hecho, sin colocar ninguna orden.", b=True)
tabla(["Qué", "Parámetros", "Dónde se lee"],
      [["Trailing stop sobre puntos de SPX",
        "Activa a +10 pts · colchón 2.5 × ATR de 2m · el ATR se congela al entrar",
        "GET /api/spx/shadow-trail"],
       ["Resultados por versión del algoritmo",
        "Agrupa por huella, excluye el P&L no confiable, marca muestra <30 trades",
        "GET /api/spx/version-stats"]],
      [2.1, 3.1, 1.8])
par("El trailing seguirá en sombra hasta tener trades propios suficientes. La razón para no "
    "activarlo ya: cambiar el TP es de impacto ALTO, y ya entraron dos cambios de impacto "
    "alto (gatillo y umbral del 90%). Si se activa un tercero ahora, al medir no se podrá "
    "saber cuál funcionó.")

h("Lo que quedó bajo sospecha y sin resolver", 1)
bullet("EL STOP TÉCNICO (Fractal 15m / POC). Fue el 39% de los cierres con pérdida media "
       "alta, pero ese número se apoya en registros cuyo P&L está mal calculado. No se puede "
       "acusar ni absolver con los datos actuales.", ROJO)
bullet("EL HISTÓRICO ESTÁ CONTAMINADO. De 62 trades direccionales, 39 tienen el P&L mal "
       "calculado por el método anterior; al menos 5 registran una pérdida MAYOR al débito "
       "pagado, algo imposible. Son irrecuperables: Tradier ya no lista esas órdenes.", ROJO)
bullet("EL PUNTO DE PARTIDA REAL son 51 trades con P&L confiable: 45.1% de win rate. "
       "El objetivo es 80%.")
bullet("LA FRECUENCIA. El mentor toma 1-2 operaciones por día; el sistema 4-6. Como sus "
       "entradas dan 89% de acierto incluso con NUESTRA salida actual, la diferencia debe "
       "estar en las operaciones que el sistema toma y él no. Es el pendiente más grande.", ROJO)

h("Advertencia sobre toda la evidencia de este documento", 1)
par("Las 14 fotos son SOLO operaciones exitosas del mentor. Que el gatillo coincida con sus "
    "aciertos confirma que se está leyendo el mismo setup —lo cual ya es valioso— pero no "
    "dice nada sobre cuántas veces dispara donde él no entra, ni qué pasa cuando el setup "
    "falla. Es sesgo de supervivencia.", c=ROJO, b=True)
par("La ventaja del gatillo de pullback sobre la lógica anterior sí está medida de forma "
    "independiente (18 días de velas reales, 7 combinaciones de TP/SL, y las dos mitades de "
    "la muestra por separado). Pero esa comparación es ENTRE VARIANTES DE SIMULACIÓN. "
    "El pullback debería ser mejor que lo anterior; no hay evidencia suficiente para "
    "afirmar que sea rentable.")

h("Qué vigilar los primeros días", 1)
tabla(["Qué", "Referencia", "Si se desvía"],
      [["Señales por día", "5-6 esperadas (antes 3.9)", "Muy por encima: apretar el 0.8 × ATR"],
       ["Primer trade bajista en vivo", "Nunca observado fuera de reconstrucción",
        "Revisarlo vela por vela"],
       ["Rechazos por score", "Desconocido",
        "Si el gatillo mejora y el score se come la mejora, revisar los pesos"],
       ["Registros del trailing en sombra", "Sin datos aún",
        "Comparar contra el cierre real de cada trade"]],
      [1.9, 2.5, 2.6])

os.makedirs(DEST, exist_ok=True)
salida = os.path.join(DEST, ARCHIVO)
try:
    d.save(salida)
    print("guardado:", salida)
except PermissionError:
    alt = salida.replace(".docx", "_nuevo.docx")
    d.save(alt)
    print("ABIERTO EN WORD, guardado como:", alt)
