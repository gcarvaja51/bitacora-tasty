# -*- coding: utf-8 -*-
"""
Antepone la pagina "Informe Gerencial" (resumen ejecutivo de UNA sola hoja) al
documento de premercado ya generado.

Uso:
    python gen_informe_gerencial.py <docx> <spec.json> <chart_escenarios.png> [--force]

Corre DESPUES de que el .docx del dia ya esta escrito y DESPUES de que
gen_escenarios_chart.py dejo el PNG en disco -- reusa ese mismo PNG, no genera
un grafico nuevo, para que la portada y el cuerpo no puedan contradecirse.

La pagina se INSERTA antes del primer parrafo del documento (el Heading 1
"Premercado SPX - <fecha>") y se cierra con un salto de pagina, de modo que el
analisis completo arranca intacto en la hoja 2. No se toca ni un parrafo del
cuerpo.

Idempotencia: si el documento ya empieza con el Informe Gerencial, aborta con
exit 3 salvo que se pase --force (que en ese caso BORRA el bloque anterior y lo
reescribe -- util cuando hay que corregir un dato despues de haberlo generado).

--- spec.json ---
{
  "fecha": "viernes 14 de agosto de 2026",
  "contexto": "Ejecutado 9:20 AM ET  |  Cierre previo 7.798,99  |  ...",
  "tesis": "Parrafo de 4-6 lineas. La conclusion, no el razonamiento.",
  "termometro": [["Regimen gamma", "POSITIVO"], ["Net GEX", "+$21,35 B"]],
  "cambios": ["Bullet corto de que cambio vs ayer", "..."],
  "escenarios": [
    {"nombre": "Neutral - Pinning en 7.800", "prob": "46%",
     "activa": "...", "invalida": "...", "target": "...", "color": "neutral"}
  ],
  "primera_hora": ["Bullet 1", "Bullet 2", "Bullet 3"],
  "riesgo": "Parrafo de 2-3 lineas con el riesgo principal del dia."
}

"color" de cada escenario: "alcista" | "bajista" | "neutral" -- solo tine el
nombre del escenario, nunca es el unico portador de informacion (la etiqueta de
texto ya dice cual es), per la regla de secondary encoding del skill dataviz.
"""
import sys
import json

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FUENTE = "Tahoma"

# Misma paleta que gen_escenarios_chart.py -- la portada y el diagrama tienen que
# leerse como una sola pieza.
VERDE = RGBColor(0x00, 0x83, 0x00)
ROJO = RGBColor(0xE3, 0x49, 0x48)
INK = RGBColor(0x0B, 0x0B, 0x0B)
INK_SEC = RGBColor(0x52, 0x51, 0x4E)
INK_MUTED = RGBColor(0x89, 0x87, 0x81)

COLOR_ESC = {"alcista": VERDE, "bajista": ROJO, "neutral": INK_SEC}

ANCHO_CHART = 3.05      # pulgadas -- columna izquierda del bloque superior
ANCHO_LATERAL = 3.85    # pulgadas -- columna derecha
ANCHO_UTIL = 7.0        # 8.5" de hoja - 0.75" x 2 de margen


def _fuente(run, size=9.5, bold=False, italic=False, color=None):
    """Tahoma explicito en cada run: el estilo base no siempre gana dentro de
    celdas de tabla (gotcha ya documentado en SKILL.md)."""
    run.font.name = FUENTE
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    # eastAsia cubre tildes/enie correctamente (mismo motivo que en el cuerpo)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), FUENTE)


def _sin_espacio(par, antes=0, despues=2):
    par.paragraph_format.space_before = Pt(antes)
    par.paragraph_format.space_after = Pt(despues)


def _sin_bordes(tabla):
    """Tabla de maquetacion: rejilla invisible."""
    tbl_pr = tabla._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def _sombrear(celda, hexcolor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    celda._tc.get_or_add_tcPr().append(shd)


def _anchos(tabla, anchos_in):
    """python-docx exige fijar el ancho celda por celda ademas de apagar autofit,
    o Word recalcula la tabla a su gusto y la maqueta se descuadra."""
    tabla.autofit = False
    for fila in tabla.rows:
        for celda, ancho in zip(fila.cells, anchos_in):
            celda.width = Inches(ancho)


def _rotulo(doc, texto, color=INK, size=10):
    """Encabezado de bloque de la portada. Parrafo normal con run en negrita --
    NO usa los estilos Heading para no ensuciar el esquema del documento."""
    p = doc.add_paragraph()
    _sin_espacio(p, antes=6, despues=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _fuente(p.add_run(texto.upper()), size=size, bold=True, color=color)
    return p


def construir(doc, spec, chart_png):
    """Crea los elementos de la portada al final del documento y devuelve la
    lista de elementos XML en orden, para moverlos despues al principio."""
    creados = []

    def marcar(*elementos):
        creados.extend(elementos)

    # --- Titulo -------------------------------------------------------------
    titulo = doc.add_heading(level=1)
    _sin_espacio(titulo, antes=0, despues=1)
    _fuente(titulo.add_run(f"Informe Gerencial — Premercado SPX · {spec['fecha']}"),
            size=11, bold=True, color=INK)
    marcar(titulo._p)

    ctx = doc.add_paragraph()
    _sin_espacio(ctx, antes=0, despues=6)
    _fuente(ctx.add_run(spec["contexto"]), size=8.5, italic=True, color=INK_MUTED)
    marcar(ctx._p)

    # --- Bloque superior: grafico | tesis + termometro + cambios ------------
    superior = doc.add_table(rows=1, cols=2)
    _sin_bordes(superior)
    _anchos(superior, [ANCHO_CHART + 0.1, ANCHO_LATERAL])
    izq, der = superior.rows[0].cells
    izq.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    der.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p_img = izq.paragraphs[0]
    _sin_espacio(p_img, antes=0, despues=0)
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(chart_png, width=Inches(ANCHO_CHART))

    # Columna derecha
    p = der.paragraphs[0]
    _sin_espacio(p, antes=0, despues=2)
    _fuente(p.add_run("TESIS DEL DÍA"), size=10, bold=True, color=INK)

    p = der.add_paragraph()
    _sin_espacio(p, antes=0, despues=6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _fuente(p.add_run(spec["tesis"]), size=9, color=INK)

    p = der.add_paragraph()
    _sin_espacio(p, antes=0, despues=2)
    _fuente(p.add_run("TERMÓMETRO DEL DÍA"), size=10, bold=True, color=INK)

    termo = der.add_table(rows=len(spec["termometro"]), cols=2)
    _sin_bordes(termo)
    _anchos(termo, [1.55, 2.25])
    for i, (clave, valor) in enumerate(spec["termometro"]):
        cp = termo.rows[i].cells[0].paragraphs[0]
        _sin_espacio(cp, antes=0, despues=0)
        # LEFT explicito: el estilo Normal del documento esta en JUSTIFY y en una
        # celda angosta eso abre huecos enormes entre palabras cuando la etiqueta
        # se parte en dos lineas ("Pivote        (Call / Wall+MVS+POC)").
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _fuente(cp.add_run(clave), size=8.5, bold=True, color=INK_SEC)
        vp = termo.rows[i].cells[1].paragraphs[0]
        _sin_espacio(vp, antes=0, despues=0)
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _fuente(vp.add_run(valor), size=8.5, color=INK)
        if i % 2 == 0:
            for c in termo.rows[i].cells:
                _sombrear(c, "F0EFEC")

    if spec.get("cambios"):
        p = der.add_paragraph()
        _sin_espacio(p, antes=6, despues=2)
        _fuente(p.add_run("LO QUE CAMBIÓ VS. AYER"), size=10, bold=True, color=INK)
        for bullet in spec["cambios"]:
            b = der.add_paragraph()
            _sin_espacio(b, antes=0, despues=1)
            b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            b.paragraph_format.left_indent = Inches(0.12)
            _fuente(b.add_run(f"· {bullet}"), size=8.5, color=INK)

    marcar(superior._tbl)

    # --- Los 3 escenarios en una linea --------------------------------------
    marcar(_rotulo(doc, "Los 3 escenarios", size=10)._p)

    cols = ["Escenario", "Prob.", "Se activa si", "Se invalida si", "Target"]
    anchos = [1.65, 0.45, 1.95, 1.45, 1.50]
    tabla = doc.add_table(rows=1 + len(spec["escenarios"]), cols=len(cols))
    tabla.style = doc.styles["Light Grid Accent 1"]
    _anchos(tabla, anchos)

    for j, titulo_col in enumerate(cols):
        cp = tabla.rows[0].cells[j].paragraphs[0]
        _sin_espacio(cp, antes=1, despues=1)
        _fuente(cp.add_run(titulo_col), size=8, bold=True, color=INK)

    for i, e in enumerate(spec["escenarios"], start=1):
        celdas = tabla.rows[i].cells
        color = COLOR_ESC.get(e.get("color", "neutral"), INK_SEC)
        valores = [e["nombre"], e["prob"], e["activa"], e["invalida"], e["target"]]
        for j, valor in enumerate(valores):
            cp = celdas[j].paragraphs[0]
            _sin_espacio(cp, antes=1, despues=1)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
            # Solo el nombre y la probabilidad van tenidos: el color acompana a
            # una etiqueta que ya dice el escenario, nunca lo sustituye.
            _fuente(cp.add_run(valor), size=8,
                    bold=(j <= 1), color=color if j <= 1 else INK)
    marcar(tabla._tbl)

    # --- Primera hora | riesgo (lado a lado) --------------------------------
    # Apiladas a lo ancho de la hoja se comian 2,5 pulgadas y empujaban la
    # portada a una segunda pagina (medido el 2026-08-14 exportando a PDF).
    # En dos columnas ocupan 1,4 y ademas se leen mejor: el "que hacer" a la
    # izquierda y el "que puede salir mal" a la derecha.
    inferior = doc.add_table(rows=1, cols=2)
    _sin_bordes(inferior)
    _anchos(inferior, [4.15, 2.85])
    col_a, col_b = inferior.rows[0].cells
    col_a.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    col_b.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = col_a.paragraphs[0]
    _sin_espacio(p, antes=0, despues=2)
    _fuente(p.add_run("PRIMERA HORA — QUÉ MIRAR"), size=10, bold=True, color=INK)
    for bullet in spec["primera_hora"]:
        b = col_a.add_paragraph()
        _sin_espacio(b, antes=0, despues=2)
        b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        b.paragraph_format.left_indent = Inches(0.12)
        _fuente(b.add_run(f"· {bullet}"), size=8.5, color=INK)

    p = col_b.paragraphs[0]
    _sin_espacio(p, antes=0, despues=2)
    _fuente(p.add_run("RIESGO PRINCIPAL"), size=10, bold=True, color=ROJO)
    p = col_b.add_paragraph()
    _sin_espacio(p, antes=0, despues=0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _fuente(p.add_run(spec["riesgo"]), size=8.5, color=INK)

    marcar(inferior._tbl)

    return creados


def es_titulo_cuerpo(el):
    """El Heading 1 original del documento, que marca donde empieza el analisis
    completo. Es el ancla de insercion y tambien el limite del bloque gerencial."""
    return (el.tag.split("}")[-1] == "p"
            and "".join(el.itertext()).strip().startswith("Premercado SPX"))


def bloque_previo(doc):
    """Elementos del Informe Gerencial ya insertado en una corrida anterior:
    desde el principio del cuerpo hasta (sin incluir) el titulo del analisis."""
    cuerpo = list(doc.element.body.iterchildren())
    if not cuerpo:
        return []
    primero = cuerpo[0]
    if primero.tag.split("}")[-1] != "p" or "Informe Gerencial" not in "".join(primero.itertext()):
        return []
    previos = []
    for el in cuerpo:
        if es_titulo_cuerpo(el):
            break
        previos.append(el)
    return previos


def main():
    if len(sys.argv) < 4:
        print("uso: gen_informe_gerencial.py <docx> <spec.json> <chart.png> [--force]")
        sys.exit(1)

    docx_path, spec_path, chart_png = sys.argv[1], sys.argv[2], sys.argv[3]
    force = "--force" in sys.argv[4:]

    with open(spec_path, encoding="utf-8") as fh:
        spec = json.load(fh)

    doc = Document(docx_path)

    previos = bloque_previo(doc)
    if previos and not force:
        print("ABORTA: el documento ya tiene Informe Gerencial. Usar --force para rehacerlo.")
        sys.exit(3)
    if previos and force:
        for el in previos:
            el.getparent().remove(el)
        print("Bloque gerencial anterior eliminado (%d elementos)." % len(previos))

    ancla = doc.element.body.find(qn("w:p"))
    if ancla is None:
        print("ABORTA: el documento no tiene parrafos.")
        sys.exit(1)

    creados = construir(doc, spec, chart_png)
    for el in creados:
        ancla.addprevious(el)

    # El corte de hoja va como "salto antes" del titulo del analisis, NO como un
    # parrafo con salto al final de la portada: ese parrafo extra no cabia en la
    # hoja 1 ya llena, se iba solo a la hoja 2 y empujaba el cuerpo a la 3,
    # dejando una pagina en blanco (visto al exportar a PDF el 2026-08-14).
    Paragraph(ancla, doc).paragraph_format.page_break_before = True

    doc.save(docx_path)
    print("OK: informe gerencial insertado (%d elementos)." % len(creados))


if __name__ == "__main__":
    main()
